"""
调试批注段落关联
用法：python3 test_parse4.py <docx文件路径>
"""
import sys
import os

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from lxml import etree

def debug_comment_ranges(file_path):
    doc = Document(file_path)
    
    print("=" * 60)
    print("调试批注段落关联")
    print("=" * 60)
    
    # 1. 打印所有段落（带索引）
    print("\n[1] 所有段落:")
    for i, para in enumerate(doc.paragraphs):
        style = para.style.name if para.style else 'Normal'
        text = para.text[:60] + "..." if len(para.text) > 60 else para.text
        marker = ""
        if 'Heading' in style:
            marker = " <<<"
        print(f"  [{i}] [{style}]{marker} {text}")
    
    # 2. 解析 document.xml 中的 commentRangeStart/End
    print("\n[2] commentRangeStart/End 位置:")
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    doc_xml = etree.fromstring(doc.part.blob)
    body = doc_xml.find('.//w:body', nsmap)
    
    para_idx = 0
    comment_ranges = {}
    
    for elem in body.iter():
        tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        
        if tag == 'p':
            # 打印段落开头文本
            texts = [t.text for t in elem.findall('.//w:t', nsmap) if t.text]
            para_text = ''.join(texts)[:40]
            print(f"  XML段落[{para_idx}]: {para_text}...")
            para_idx += 1
            
        elif tag == 'commentRangeStart':
            cid = elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
            print(f"    >>> commentRangeStart id={cid} (当前段落={para_idx})")
            if cid and cid not in comment_ranges:
                comment_ranges[cid] = {'start_para': para_idx, 'end_para': para_idx}
                
        elif tag == 'commentRangeEnd':
            cid = elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
            print(f"    >>> commentRangeEnd id={cid} (当前段落={para_idx})")
            if cid and cid in comment_ranges:
                comment_ranges[cid]['end_para'] = para_idx
    
    # 3. 获取批注文本
    print("\n[3] 批注关联结果:")
    comments = []
    for rel in doc.part.rels.values():
        reltype = str(rel.reltype)
        if reltype.endswith('/comments') and 'commentsExtended' not in reltype:
            comments_part = rel.target_part
            root = etree.fromstring(comments_part.blob)
            for comment in root.findall('.//w:comment', nsmap):
                cid = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                texts = []
                for p in comment.findall('.//w:p', nsmap):
                    t = [x.text for x in p.findall('.//w:t', nsmap) if x.text]
                    texts.append(''.join(t))
                comments.append({'id': cid, 'text': '\n'.join(texts).strip()[:50]})
            break
    
    comments_by_id = {c['id']: c['text'] for c in comments}
    
    for cid, rng in comment_ranges.items():
        text = comments_by_id.get(cid, '(未知)')
        print(f"  批注 ID={cid}: 段落[{rng['start_para']}] ~ 段落[{rng['end_para']}]")
        print(f"    内容: {text}...")
        
        # 显示关联的段落文本
        for pi in range(rng['start_para'], rng['end_para'] + 1):
            if pi < len(doc.paragraphs):
                p = doc.paragraphs[pi]
                print(f"    -> 段落[{pi}]: {p.text[:40]}...")

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("用法: python3 test_parse4.py <docx文件路径>")
        sys.exit(1)
    debug_comment_ranges(sys.argv[1])
