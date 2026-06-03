"""
Word 结构优先解析器
专门处理 Word 文档，精确提取：
- 书名（Heading 1）
- 作者、国籍、版本（书名后的 Normal 段落）
- 章（Heading 2）
- 节（Heading 3）
- 批注（对标题的批注=小结，对正文的批注=点评）
"""
from docx import Document
from lxml import etree
import re


class WordStructureParser:
    """Word 文档结构解析器"""
    
    # Word 命名空间
    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    
    def __init__(self, file_path):
        self.doc = Document(file_path)
        self.paragraphs = list(self.doc.paragraphs)
        self.nsmap = {'w': self.W_NS}
        
    def parse(self):
        """解析整个文档"""
        # 1. 提取元信息
        meta = self._extract_metadata()
        
        # 2. 提取批注信息
        comments = self._extract_comments()
        comment_ranges = self._extract_comment_ranges()
        
        # 3. 构建章节结构
        chapters, sections = self._build_structure(comments, comment_ranges)
        
        return {
            'title': meta['title'],
            'author': meta['author'],
            'author_nationality': meta['nationality'],
            'version': meta['version'],
            'chapters': chapters,
            'sections': sections
        }
    
    def _extract_metadata(self):
        """
        提取文档元信息：
        - 书名：Heading 1
        - 作者/国籍：书名后的第一个 Normal 段落
        - 版本：书名后的第二个 Normal 段落
        """
        meta = {
            'title': '未命名',
            'author': '',
            'nationality': '',
            'version': ''
        }
        
        # 查找第一个 Heading 1 作为书名
        for para in self.paragraphs:
            style = para.style.name if para.style else 'Normal'
            if style == 'Heading 1':
                meta['title'] = para.text.strip()
                break
        
        # 查找书名后的段落作为作者/版本
        found_title = False
        normal_count = 0
        for para in self.paragraphs:
            style = para.style.name if para.style else 'Normal'
            text = para.text.strip()
            
            if style == 'Heading 1':
                found_title = True
                continue
            
            if not found_title:
                continue
            
            # 支持自定义样式 [作者] 和 [版本]
            if style == '[作者]' or style == '作者':
                # 只取第一个 [作者] 段落
                if meta.get('author'):
                    continue
                if '「' in text and '」' in text:
                    parts = text.split('「')
                    meta['author'] = parts[0].strip()
                    meta['nationality'] = parts[1].split('」')[0].strip()
                elif '[' in text and ']' in text:
                    parts = text.split('[')
                    meta['author'] = parts[0].strip()
                    meta['nationality'] = parts[1].split(']')[0].strip()
                elif '（' in text and '）' in text:
                    parts = text.split('（')
                    meta['author'] = parts[0].strip()
                    meta['nationality'] = parts[1].split('）')[0].strip()
                else:
                    # 如果作者行没有格式，尝试从文本中解析
                    meta['author'] = text
                    # 尝试匹配 高尔基「苏联」 格式
                    import re
                    match = re.search(r'(.+?)「(.+?)」', text)
                    if match:
                        meta['author'] = match.group(1).strip()
                        meta['nationality'] = match.group(2).strip()
                    match = re.search(r'(.+?)\[(.+?)\]', text)
                    if match:
                        meta['author'] = match.group(1).strip()
                        meta['nationality'] = match.group(2).strip()
                    match = re.search(r'(.+?)（(.+?)）', text)
                    if match:
                        meta['author'] = match.group(1).strip()
                        meta['nationality'] = match.group(2).strip()
                continue
            
            if style == '[版本]' or style == '版本':
                # 只取第一个 [版本] 段落
                if meta.get('version'):
                    continue
                meta['version'] = text
                continue
            
            # 处理 Normal 样式作为备用（仅当未通过 [作者]/[版本] 样式获取时）
            if style == 'Normal' and text:
                normal_count += 1
                if normal_count == 1 and not meta.get('author'):
                    # 第一行：作者「国籍」
                    if '「' in text and '」' in text:
                        parts = text.split('「')
                        meta['author'] = parts[0].strip()
                        meta['nationality'] = parts[1].split('」')[0].strip()
                    elif '[' in text and ']' in text:
                        parts = text.split('[')
                        meta['author'] = parts[0].strip()
                        meta['nationality'] = parts[1].split(']')[0].strip()
                    elif '（' in text and '）' in text:
                        parts = text.split('（')
                        meta['author'] = parts[0].strip()
                        meta['nationality'] = parts[1].split('）')[0].strip()
                    else:
                        meta['author'] = text
                elif normal_count == 2 and not meta.get('version'):
                    # 第二行：版本
                    meta['version'] = text
                    break
        
        return meta
    
    def _extract_comments(self):
        """提取所有批注内容"""
        comments = {}
        
        for rel in self.doc.part.rels.values():
            reltype = str(rel.reltype)
            if reltype.endswith('/comments') and 'commentsExtended' not in reltype:
                comments_part = rel.target_part
                root = etree.fromstring(comments_part.blob)
                
                for comment in root.findall('.//w:comment', self.nsmap):
                    cid = comment.get(f'{{{self.W_NS}}}id')
                    author = comment.get(f'{{{self.W_NS}}}author', '')
                    
                    texts = []
                    for p in comment.findall('.//w:p', self.nsmap):
                        for t in p.findall('.//w:t', self.nsmap):
                            if t.text:
                                texts.append(t.text)
                    text = ''.join(texts).strip()
                    
                    comments[cid] = {
                        'id': cid,
                        'author': author,
                        'text': text
                    }
        
        return comments
    
    def _extract_comment_ranges(self):
        """提取批注引用范围"""
        comment_ranges = {}
        
        for para_idx, para in enumerate(self.paragraphs):
            para_element = para._element
            char_offset = 0
            
            for child in para_element:
                tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                
                if tag == 'commentRangeStart':
                    cid = child.get(f'{{{self.W_NS}}}id')
                    if cid and cid not in comment_ranges:
                        comment_ranges[cid] = {
                            'start_para': para_idx,
                            'start_char': char_offset,
                            'end_para': para_idx,
                            'end_char': char_offset
                        }
                
                elif tag == 'commentRangeEnd':
                    cid = child.get(f'{{{self.W_NS}}}id')
                    if cid and cid in comment_ranges:
                        comment_ranges[cid]['end_para'] = para_idx
                        comment_ranges[cid]['end_char'] = char_offset
                
                elif tag == 'r':
                    # 计算该 run 中的所有字符（包括文本和TAB等特殊字符）
                    for elem in child.iter():
                        elem_tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
                        if elem_tag == 't' and elem.text:
                            char_offset += len(elem.text)
                        elif elem_tag == 'tab':
                            # TAB 字符也算一个字符
                            char_offset += 1
        
        return comment_ranges
    
    def _extract_number(self, text):
        """从标题文本中提取序号，如 '第一章' -> 1, '1. 标题' -> 1, '第3节' -> 3, '第一篇' -> 1"""
        # 匹配 "第X章"、"第X节"、"第X回"、"第X篇" 等中文序号
        match = re.search(r'第\s*(\d+)\s*[章节回篇]', text)
        if match:
            return int(match.group(1))
        # 匹配 "1."、"1、"、"1 " 开头的数字序号
        match = re.match(r'^\s*(\d+)\s*[.．、\s]', text)
        if match:
            return int(match.group(1))
        # 匹配纯数字开头
        match = re.match(r'^\s*(\d+)\s*', text)
        if match:
            return int(match.group(1))
        return None

    def _get_style_numPr(self, para):
        """从段落样式中获取 numPr"""
        try:
            p_elem = para._element
            pStyle = p_elem.find(f'{{{self.W_NS}}}pPr/{{{self.W_NS}}}pStyle')
            if pStyle is not None:
                style_val = pStyle.get(f'{{{self.W_NS}}}val')
                if style_val:
                    for style in self.doc.styles:
                        if style.style_id == style_val:
                            style_elem = style.element
                            style_numPr = style_elem.find(f'.//{{{self.W_NS}}}numPr')
                            if style_numPr is not None:
                                return style_numPr
        except Exception as e:
            print(f"[Parser] 样式numPr查找失败: {e}")
        return None

    def _get_auto_number_text(self, para):
        """从 Word XML 中提取段落的自动编号格式
        返回: (lvlText, numFmt, start) 或 None
        """
        try:
            p_elem = para._element
            
            # 获取段落直接属性中的 numPr
            numPr = p_elem.find(f'{{{self.W_NS}}}pPr/{{{self.W_NS}}}numPr')
            
            # numId=0 表示明确移除编号，直接返回 None
            if numPr is not None:
                numId_elem = numPr.find(f'{{{self.W_NS}}}numId')
                if numId_elem is not None:
                    num_id_val = numId_elem.get(f'{{{self.W_NS}}}val')
                    if num_id_val == '0':
                        return None
            
            # 如果段落没有 numPr，尝试从样式获取
            effective_numPr = numPr
            if effective_numPr is None:
                effective_numPr = self._get_style_numPr(para)
            
            if effective_numPr is None:
                return None
            
            # 获取 ilvl（可能不存在，默认0）
            ilvl = effective_numPr.find(f'{{{self.W_NS}}}ilvl')
            level = int(ilvl.get(f'{{{self.W_NS}}}val', '0')) if ilvl is not None else 0
            # 获取 numId（可能为0或不存在）
            numId = effective_numPr.find(f'{{{self.W_NS}}}numId')
            num_id_val = numId.get(f'{{{self.W_NS}}}val') if numId is not None else None
            
            fmt = self._get_numbering_format(num_id_val, level)
            if fmt:
                print(f"[Parser] 获取到编号格式: lvlText='{fmt[0]}', numFmt={fmt[1]}, start={fmt[2]}")
                return fmt
            return None
        except Exception as e:
            print(f"[Parser] 获取自动编号格式失败: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def _get_numbering_xml(self):
        """获取 numbering.xml 的根元素"""
        try:
            numbering_part = self.doc.part.numbering_part
            return numbering_part._element
        except Exception as e:
            print(f"[Parser] 获取numbering.xml失败: {e}")
            return None
    
    def _get_numbering_format(self, num_id_val, level):
        """
        从 numbering.xml 获取编号格式信息（不计算具体编号值）
        返回: (lvlText, numFmt, start) 或 None
        """
        numbering_elem = self._get_numbering_xml()
        if numbering_elem is None:
            return None
        
        W = self.W_NS
        
        # 方法1：通过 w:num -> w:abstractNumId 查找
        if num_id_val and str(num_id_val) != '0':
            for num in numbering_elem.findall(f'{{{W}}}num'):
                found_num_id = num.get(f'{{{W}}}numId')
                if found_num_id == str(num_id_val):
                    abstract_num_id_ref = num.find(f'{{{W}}}abstractNumId')
                    if abstract_num_id_ref is not None:
                        abstract_num_id = abstract_num_id_ref.get(f'{{{W}}}val')
                        result = self._find_abstract_num_level_format(numbering_elem, abstract_num_id, level)
                        if result:
                            return result
                    break
        
        # 方法2：直接遍历 w:abstractNum
        return self._find_numbering_format_by_style(numbering_elem, level)
    
    def _find_abstract_num_level_format(self, numbering_elem, abstract_num_id, level):
        """通过 abstractNumId 和 level 查找编号格式"""
        W = self.W_NS
        for an in numbering_elem.findall(f'{{{W}}}abstractNum'):
            found_id = an.get(f'{{{W}}}abstractNumId')
            if found_id == str(abstract_num_id):
                for lvl in an.findall(f'{{{W}}}lvl'):
                    found_ilvl = lvl.get(f'{{{W}}}ilvl')
                    if found_ilvl is not None and int(found_ilvl) == level:
                        return self._extract_level_format(lvl)
        return None
    
    def _find_numbering_format_by_style(self, numbering_elem, level):
        """遍历所有 abstractNum，查找匹配当前 level 的编号格式"""
        W = self.W_NS
        for an in numbering_elem.findall(f'{{{W}}}abstractNum'):
            for lvl in an.findall(f'{{{W}}}lvl'):
                found_ilvl = lvl.get(f'{{{W}}}ilvl')
                if found_ilvl is not None and int(found_ilvl) == level:
                    result = self._extract_level_format(lvl)
                    if result:
                        return result
        return None
    
    def _extract_level_format(self, lvl_elem):
        """从 w:lvl 元素提取编号格式"""
        W = self.W_NS
        numFmt_elem = lvl_elem.find(f'{{{W}}}numFmt')
        lvlText_elem = lvl_elem.find(f'{{{W}}}lvlText')
        start_elem = lvl_elem.find(f'{{{W}}}start')
        
        numFmt = numFmt_elem.get(f'{{{W}}}val', 'decimal') if numFmt_elem is not None else 'decimal'
        lvlText = lvlText_elem.get(f'{{{W}}}val', '%1.') if lvlText_elem is not None else '%1.'
        start = int(start_elem.get(f'{{{W}}}val', '1')) if start_elem is not None else 1
        
        return (lvlText, numFmt, start)
    
    def _format_number_text(self, lvlText, numFmt, num_val):
        """根据格式和数值生成编号文本"""
        import re
        result = lvlText
        placeholders = re.findall(r'%(\d+)', lvlText)
        for ph in placeholders:
            ph_num = int(ph)
            if ph_num == 1:
                if numFmt == 'chineseCounting':
                    val = self._num_to_chinese(num_val)
                else:
                    val = str(num_val)
                result = result.replace(f'%{ph_num}', val)
        return result

    def _num_to_chinese(self, num):
        """将数字转换为中文数字"""
        chinese_nums = ['', '一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
        if num <= 10:
            return chinese_nums[num]
        elif num < 20:
            return '十' + chinese_nums[num - 10]
        elif num < 100:
            tens = num // 10
            ones = num % 10
            if ones == 0:
                return chinese_nums[tens] + '十'
            else:
                return chinese_nums[tens] + '十' + chinese_nums[ones]
        return str(num)


    def _build_structure(self, comments, comment_ranges):
        """
        构建章节结构：
        - Heading 2 = 章
        - Heading 3 = 节（阅读单元）
        - 如果只有 Heading 2 没有 Heading 3，则 Heading 2 作为节
        - 标题 = Word 自动编号（如有）+ para.text，与 Word 显示绝对一致
        - chapter_number / section_number 仅用于内部排序
        """
        chapters = []
        sections = []
        chapter_number = 0
        section_number = 0
        chapter_counter = 0   # 章自动编号计数器（仅对带编号的Heading 2递增）
        section_counter = 0   # 节自动编号计数器（每章重置）
        current_chapter_title = None
        
        # 先检查是否有 Heading 3
        has_h3 = self._has_heading3()
        print(f"[Parser] 文档是否有 Heading 3: {has_h3}")
        
        for para_idx, para in enumerate(self.paragraphs):
            style = para.style.name if para.style else 'Normal'
            text = para.text.strip()
            
            if not text:
                continue
            
            if style == 'Heading 1':
                # 书名，跳过
                continue
            
            elif style == 'Heading 2':
                # 新章节开始，重置节序号和节编号计数器
                section_number = 0
                section_counter = 0
                chapter_number += 1

                # 解析 Word 自动编号，与 para.text 拼接成完整标题
                num_fmt = self._get_auto_number_text(para)
                if num_fmt:
                    lvlText, numFmt, start = num_fmt
                    chapter_counter += 1
                    auto_num_text = self._format_number_text(lvlText, numFmt, chapter_counter)
                    title = f"{auto_num_text} {text}" if text else auto_num_text
                    print(f"[Parser] 章编号: {auto_num_text} -> 标题: {title}")
                else:
                    title = text

                if has_h3:
                    # 有 Heading 3 时，Heading 2 是章
                    current_chapter_title = title
                    chapters.append({
                        'chapter_number': chapter_number,
                        'title': title
                    })
                    print(f"[Parser] 章 {chapter_number}: {title}")
                else:
                    # 没有 Heading 3 时，Heading 2 是节
                    section = self._create_section(
                        section_number=chapter_number,
                        chapter_number=None,
                        title=title,
                        para_idx=para_idx,
                        comments=comments,
                        comment_ranges=comment_ranges
                    )
                    sections.append(section)
            
            elif style == 'Heading 3':
                # 节标题
                section_number += 1

                # 解析 Word 自动编号，与 para.text 拼接成完整标题
                num_fmt = self._get_auto_number_text(para)
                if num_fmt:
                    lvlText, numFmt, start = num_fmt
                    section_counter += 1
                    auto_num_text = self._format_number_text(lvlText, numFmt, section_counter)
                    title = f"{auto_num_text} {text}" if text else auto_num_text
                    print(f"[Parser] 节编号: {auto_num_text} -> 标题: {title}")
                else:
                    title = text

                section = self._create_section(
                    section_number=section_number,
                    chapter_number=chapter_number if has_h3 else None,
                    title=title,
                    para_idx=para_idx,
                    comments=comments,
                    comment_ranges=comment_ranges
                )
                sections.append(section)
        
        return chapters, sections
    
    def _has_heading3(self):
        """检查文档是否有 Heading 3 样式"""
        for para in self.paragraphs:
            style = para.style.name if para.style else 'Normal'
            if style == 'Heading 3':
                return True
        return False
    
    def _create_section(self, section_number, chapter_number, title, para_idx, comments, comment_ranges):
        """创建一个节"""
        section = {
            'section_number': section_number,
            'chapter_number': chapter_number,
            'title': title,
            'content': '',
            'summary': None,
            'annotations': []
        }
        
        print(f"[Parser] 节 {section_number}: {title}")
        
        # 检查节标题是否有批注（小结）
        for cid, rng in comment_ranges.items():
            if rng['start_para'] == para_idx and cid in comments:
                section['summary'] = comments[cid]['text']
                print(f"[Parser]   小结: {section['summary'][:40]}...")
        
        # 收集后续正文段落（不过滤空段落，与前端一致）
        content_start = None
        for i in range(para_idx + 1, len(self.paragraphs)):
            para = self.paragraphs[i]
            style = para.style.name if para.style else 'Normal'
            text = para.text or ''  # 空段落也保留
            
            if style in ['Heading 1', 'Heading 2', 'Heading 3']:
                break
            
            if section['content']:
                section['content'] += '\n' + text
            else:
                section['content'] = text
                content_start = i
        
        # 收集正文批注（点评）
        if content_start is not None:
            content_char_offset = 0
            for i in range(content_start, len(self.paragraphs)):
                para = self.paragraphs[i]
                style = para.style.name if para.style else 'Normal'
                
                if style in ['Heading 1', 'Heading 2', 'Heading 3']:
                    break
                
                # 使用与 _extract_comment_ranges 一致的字符计算方式
                para_char_count = self._count_para_chars(para)
                
                for cid, rng in comment_ranges.items():
                    if rng['start_para'] == i and cid in comments:
                        original_text = self._get_original_text(rng)
                        abs_start = content_char_offset + rng['start_char']
                        abs_end = content_char_offset + rng['end_char']
                        
                        annotation = {
                            'original_text': original_text,
                            'comment': comments[cid]['text'],
                            'start_char': abs_start,
                            'end_char': abs_end
                        }
                        section['annotations'].append(annotation)
                        print(f"[Parser]   点评: \"{original_text[:30]}...\" (pos:{abs_start}-{abs_end})")
                
                content_char_offset += para_char_count
        
        return section
    
    def _count_para_chars(self, para):
        """计算段落中的字符数（与 _extract_comment_ranges 中的计算方式一致）"""
        char_count = 0
        para_element = para._element
        for child in para_element:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag == 'r':
                for elem in child.iter():
                    elem_tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
                    if elem_tag == 't' and elem.text:
                        char_count += len(elem.text)
                    elif elem_tag == 'tab':
                        char_count += 1
        return char_count
    
    def _get_original_text(self, rng):
        """根据批注范围提取引用的原文"""
        if rng['start_para'] == rng['end_para']:
            para_text = self.paragraphs[rng['start_para']].text
            return para_text[rng['start_char']:rng['end_char']]
        else:
            para_text = self.paragraphs[rng['start_para']].text
            return para_text[rng['start_char']:]


def parse_file(file_path):
    """解析 Word 文件的入口函数"""
    parser = WordStructureParser(file_path)
    return parser.parse()


if __name__ == '__main__':
    import sys
    if len(sys.argv) < 2:
        print("用法: python3 word_parser.py <docx文件路径>")
        sys.exit(1)
    
    result = parse_file(sys.argv[1])
    
    print("=" * 60)
    print(f"书名: {result['title']}")
    print(f"作者: {result['author']}")
    print(f"国籍: {result['author_nationality']}")
    print(f"版本: {result['version']}")
    print(f"章数: {len(result['chapters'])}")
    print(f"节数: {len(result['sections'])}")
    print("=" * 60)
    
    for ch in result['chapters'][:3]:
        print(f"\n章 {ch['chapter_number']}: {ch['title']}")
        ch_sections = [s for s in result['sections'] if s.get('chapter_number') == ch['chapter_number']]
        for sec in ch_sections:
            print(f"  节 {sec['section_number']}: {sec['title']}")
            if sec['summary']:
                print(f"    小结: {sec['summary'][:30]}...")
    
    if not result['chapters']:
        print("\n无章结构，直接显示节:")
        for sec in result['sections'][:3]:
            print(f"\n节 {sec['section_number']}: {sec['title']}")
            if sec['summary']:
                print(f"  小结: {sec['summary'][:30]}...")
            print(f"  点评: {len(sec['annotations'])} 条")
