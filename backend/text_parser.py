"""
悦读小将 - 文本解析模块
将 TXT / DOCX 文件解析为章节和小节

Word 文件结构约定：
  - Heading 1: 书名（第一个）
  - Normal: 作者「国籍」（第二个段落）
  - Normal: 版本信息（第三个段落）
  - Heading 1: 章标题（后续出现的，可选）
  - Heading 2: 节标题（第X节 节名）
  - Normal: 节正文内容
"""

import re
import os


def parse_file(file_path):
    """
    根据文件扩展名自动选择解析方式
    支持 .txt 和 .docx 格式

    返回字典：
    {
        'title': 书名,
        'author': 作者,
        'author_nationality': 国籍,
        'version': 版本,
        'chapters': [{'chapter_number': 1, 'title': '...', 'sections': [...]}],
        'sections': [{'section_number': 1, 'title': '...', 'content': '...', 'chapter_number': None}]
    }
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.docx':
        return parse_docx_file(file_path)
    else:
        return parse_txt_file(file_path)


def parse_docx_file(file_path):
    """
    解析 Word (.docx) 文件

    结构约定：
    - 第1个 Heading 1 = 书名
    - 第2个段落(Normal) = 作者「国籍」
    - 第3个段落(Normal) = 版本
    - 后续 Heading 1 = 章标题（可选）
    - Heading 2 = 节标题
    - Normal = 节正文
    """
    from docx import Document

    doc = Document(file_path)

    # 收集所有段落（带样式信息）
    all_paras = []
    for para in doc.paragraphs:
        text = para.text.strip()
        style = para.style.name if para.style else 'Normal'
        all_paras.append({'text': text, 'style': style})

    # 收集批注（comments）
    # Word批注结构：对某一节标题的批注 = 小结；对节正文的批注 = 点评
    # commentRangeStart/End 是 w:p 的直接子元素，记录精确字符偏移
    comments = []
    try:
        comments_part = None
        
        # 遍历所有关系，找到真正的 comments（跳过 commentsExtended）
        for rel in doc.part.rels.values():
            reltype = str(rel.reltype)
            if reltype.endswith('/comments') and 'commentsExtended' not in reltype and 'commentsEx' not in reltype:
                comments_part = rel.target_part
                print(f"[Parser] 找到批注部分: {reltype}")
                break
        
        if not comments_part:
            for rel in doc.part.rels.values():
                reltype = str(rel.reltype).lower()
                if 'comments' in reltype and 'extended' not in reltype and 'ex' not in reltype:
                    comments_part = rel.target_part
                    print(f"[Parser] 找到批注部分(备选): {reltype}")
                    break
        
        if comments_part:
            from lxml import etree
            nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            root = etree.fromstring(comments_part.blob)
            
            comment_elems = root.findall('.//w:comment', nsmap)
            print(f"[Parser] 找到 {len(comment_elems)} 个批注元素")
            
            for comment in comment_elems:
                comment_id = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                author = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author', '')
                comment_texts = []
                for p in comment.findall('.//w:p', nsmap):
                    texts = [t.text for t in p.findall('.//w:t', nsmap) if t.text]
                    comment_texts.append(''.join(texts))
                text = '\n'.join(comment_texts).strip()
                comments.append({
                    'id': comment_id,
                    'author': author,
                    'text': text
                })
                print(f"[Parser]   批注 ID={comment_id}: {text[:50]}...")
            print(f"[Parser] 共解析到 {len(comments)} 条批注")
        else:
            print("[Parser] 未找到标准 comments 部分，跳过批注解析")
    except Exception as e:
        print(f"[Parser] 批注解析警告: {e}")
        import traceback
        traceback.print_exc()

    # 构建精确批注范围映射
    # commentRangeStart/End 是 w:p 的直接子元素，有精确字符偏移
    para_comments = {}  # paragraph_index -> [comment]
    try:
        from lxml import etree
        nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        doc_xml = etree.fromstring(doc.part.blob)
        body = doc_xml.find('.//w:body', nsmap)
        
        comments_by_id = {c['id']: c for c in comments}
        comment_ranges = {}  # comment_id -> {start_para, start_char, end_para, end_char, text}
        
        # 遍历每个段落（使用 doc.paragraphs 保持和 all_paras 一致）
        # 通过 para._element 访问 XML 元素来查找 commentRangeStart/End
        for para_idx, para in enumerate(doc.paragraphs):
            para_element = para._element
            
            # 计算段落内每个子元素的字符偏移
            char_offset = 0
            for child in para_element:
                ctag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                
                if ctag == 'commentRangeStart':
                    cid = child.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                    if cid and cid not in comment_ranges:
                        comment_ranges[cid] = {
                            'start_para': para_idx,
                            'start_char': char_offset,
                            'end_para': para_idx,
                            'end_char': char_offset
                        }
                
                elif ctag == 'commentRangeEnd':
                    cid = child.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                    if cid and cid in comment_ranges:
                        comment_ranges[cid]['end_para'] = para_idx
                        comment_ranges[cid]['end_char'] = char_offset
                
                elif ctag == 'r':
                    # 累加 run 中文本的长度
                    for t in child.findall('.//w:t', nsmap):
                        if t.text:
                            char_offset += len(t.text)
        
        # 关联批注文本，并映射到段落
        for cid, rng in comment_ranges.items():
            if cid in comments_by_id:
                rng['text'] = comments_by_id[cid]['text']
                rng['author'] = comments_by_id[cid]['author']
                # 将批注关联到 start_para 和 end_para 之间的所有段落
                for pi in range(rng['start_para'], rng['end_para'] + 1):
                    if pi not in para_comments:
                        para_comments[pi] = []
                    para_comments[pi].append(rng)
        
        print(f"[Parser] 批注范围映射完成: {len(comment_ranges)} 条")
    except Exception as e:
        print(f"[Parser] 批注范围映射警告: {e}")
        import traceback
        traceback.print_exc()

    if not all_paras:
        return _empty_result(file_path)

    # === 1. 提取元信息 ===
    title = ''
    author = ''
    author_nationality = ''
    version = ''
    meta_end = 0  # 元信息结束位置

    # 第一个 Heading 1 是书名
    for i, p in enumerate(all_paras):
        if 'Heading 1' in p['style'] and p['text']:
            title = p['text']
            meta_end = i + 1
            break

    # 紧跟书名后的 Normal 段落：作者「国籍」
    for i in range(meta_end, min(meta_end + 2, len(all_paras))):
        p = all_paras[i]
        if 'Heading' in p['style']:
            break
        if not p['text']:
            continue
        # 匹配 "作者「国籍」" 或 "作者[国籍]" 格式
        m = re.match(r'^(.+?)[\u300c\u300d\[\【]([^\u300d\]\】]+)[\u300d\]\】]', p['text'])
        if m:
            author = m.group(1).strip()
            author_nationality = m.group(2).strip()
            meta_end = i + 1
            break
        else:
            # 没有国籍标记，尝试作为版本信息
            if not version:
                version = p['text']
                meta_end = i + 1

    # 再下一段 Normal 可能是版本信息
    for i in range(meta_end, min(meta_end + 1, len(all_paras))):
        p = all_paras[i]
        if 'Heading' in p['style']:
            break
        if not p['text']:
            continue
        if not version:
            version = p['text']
            meta_end = i + 1
            break

    # === 2. 提取章节和节 ===
    chapters = []       # 章列表
    sections = []       # 节列表（扁平）
    current_chapter = None
    current_section = None
    section_number = 0

    for i in range(meta_end, len(all_paras)):
        p = all_paras[i]
        text = p['text']
        style = p['style']

        if not text:
            continue

        # Heading 1 = 章标题（书名之后的）
        if 'Heading 1' in style and i > 0:
            # 保存当前未完成的节
            if current_section and current_section.get('content'):
                sections.append(current_section)
                current_section = None
            current_chapter = {
                'chapter_number': len(chapters) + 1,
                'title': text
            }
            chapters.append(current_chapter)
            continue

        # Heading 2 = 节标题
        if 'Heading 2' in style:
            # 保存当前未完成的节
            if current_section and current_section.get('content'):
                sections.append(current_section)

            section_number += 1
            # 提取节名（"第X节  节名" 格式）
            sec_title = text
            m = re.match(r'^第\d+节\s+(.+)', text)
            if m:
                sec_title = m.group(1).strip()

            current_section = {
                'section_number': section_number,
                'title': sec_title,
                'content': '',
                'chapter_number': current_chapter['chapter_number'] if current_chapter else 1,
                'summary': '',      # 从批注中提取的小结
                'annotations': []   # 从批注中提取的点评
            }

            # 检查该节标题段落是否有批注 -> 作为小结
            if i in para_comments:
                for cm in para_comments[i]:
                    current_section['summary'] = cm.get('text', '')
                    print(f"[Parser] 节{section_number} 标题批注作为小结: {cm.get('text', '')[:50]}...")

            # 记录当前节在 all_paras 中的起始索引，用于后续关联正文批注
            current_section['_para_start'] = i
            continue

        # Normal = 正文内容
        if current_section is not None:
            # 记录段落在 content 中的起始位置（在拼接之前）
            para_start_in_content = len(current_section['content']) if current_section['content'] else 0
            # 加上换行符的长度（如果不是第一个段落）
            if current_section['content']:
                para_start_in_content += 1  # '\n' 占1个字符
            
            if current_section['content']:
                current_section['content'] += '\n' + text
            else:
                current_section['content'] = text

            # 检查该正文段落是否有批注
            if i in para_comments:
                for cm in para_comments[i]:
                    # 使用精确字符偏移提取原文
                    start_char_in_para = cm.get('start_char', 0)
                    end_char_in_para = cm.get('end_char', len(text))
                    
                    # 如果批注跨段落，且当前是起始段落，提取从 start_char 到段落末尾
                    if cm.get('start_para') != cm.get('end_para'):
                        if i == cm.get('start_para'):
                            annotated_text = text[start_char_in_para:]
                        elif i == cm.get('end_para'):
                            annotated_text = text[:end_char_in_para]
                        else:
                            annotated_text = text
                    else:
                        # 同一段落内的批注，精确提取
                        annotated_text = text[start_char_in_para:end_char_in_para]
                    
                    if not annotated_text:
                        annotated_text = text
                    
                    # 计算在节内容中的绝对字符位置
                    # para_start_in_content 是当前段落在 content 中的起始位置
                    abs_start = para_start_in_content + start_char_in_para
                    abs_end = para_start_in_content + end_char_in_para
                    
                    current_section['annotations'].append({
                        'original_text': annotated_text,
                        'comment': cm.get('text', ''),
                        'start_char': abs_start,
                        'end_char': abs_end
                    })
                    print(f"[Parser] 节{section_number} 点评: 原文=\"{annotated_text[:30]}...\"")

    # 保存最后一个节
    if current_section and current_section.get('content'):
        sections.append(current_section)

    return {
        'title': title or get_book_title(file_path),
        'author': author,
        'author_nationality': author_nationality,
        'version': version,
        'chapters': chapters,
        'sections': sections
    }


def parse_txt_file(file_path):
    """
    解析 TXT 文件（兼容旧格式）
    返回与 parse_docx_file 相同结构的字典
    """
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    content = clean_text(content)
    paragraphs = [p.strip() for p in content.split('\n') if p.strip()]

    if not paragraphs:
        return _empty_result(file_path)

    # 尝试按章节模式解析
    chapter_patterns = [
        r'^第[一二三四五六七八九十百千零]+章',
        r'^第\d+章',
        r'^Chapter\s+\d+',
    ]

    chapters = []
    current_chapter = {'title': '前言', 'paragraphs': []}

    for para in paragraphs:
        is_chapter = False
        for pattern in chapter_patterns:
            if re.match(pattern, para, re.IGNORECASE):
                if current_chapter['paragraphs']:
                    chapters.append(current_chapter)
                current_chapter = {'title': para, 'paragraphs': []}
                is_chapter = True
                break
        if not is_chapter:
            current_chapter['paragraphs'].append(para)

    if current_chapter['paragraphs']:
        chapters.append(current_chapter)

    # 将章节分割为小节
    sections = []
    section_number = 0
    for ch_idx, chapter in enumerate(chapters):
        chapter_text = '\n'.join(chapter['paragraphs'])
        chapter_sections = split_into_sections(chapter_text, section_number + 1)
        for sec in chapter_sections:
            sec['chapter_number'] = ch_idx + 1
            sec['title'] = ''
            sections.append(sec)
        section_number += len(chapter_sections)

    return {
        'title': get_book_title(file_path),
        'author': '',
        'author_nationality': '',
        'version': '',
        'chapters': [{'chapter_number': i + 1, 'title': ch['title']} for i, ch in enumerate(chapters)],
        'sections': sections
    }


def _empty_result(file_path):
    """返回空结果"""
    return {
        'title': get_book_title(file_path),
        'author': '',
        'author_nationality': '',
        'version': '',
        'chapters': [],
        'sections': []
    }


def clean_text(text):
    """清理文本"""
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    text = re.sub(r' +', ' ', text)
    text = text.replace('"', '"').replace('"', '"')
    text = text.replace("'", "'").replace("'", "'")
    return text.strip()


def split_into_sections(text, start_number=1):
    """将文本分割为小节，每节 300-800 字"""
    sections = []
    paragraphs = text.split('\n')

    current_section = []
    current_length = 0
    section_number = start_number

    MIN_SECTION_LENGTH = 300
    MAX_SECTION_LENGTH = 800

    for para in paragraphs:
        para_length = len(para)
        if not current_section:
            current_section.append(para)
            current_length = para_length
            continue
        if current_length + para_length > MAX_SECTION_LENGTH:
            if current_length >= MIN_SECTION_LENGTH:
                sections.append({
                    'section_number': section_number,
                    'content': '\n'.join(current_section),
                    'length': current_length
                })
                section_number += 1
                current_section = [para]
                current_length = para_length
            else:
                current_section.append(para)
                current_length += para_length
        else:
            current_section.append(para)
            current_length += para_length

    if current_section:
        sections.append({
            'section_number': section_number,
            'content': '\n'.join(current_section),
            'length': current_length
        })

    return sections


def get_book_title(file_path):
    """从文件名获取书名"""
    filename = os.path.basename(file_path)
    title = os.path.splitext(filename)[0]
    return title


if __name__ == '__main__':
    # 测试
    test_path = '/workspace/.uploads/69ee5b33-6581-4cbe-8448-8830bae8bc8b_童年-高尔基-导入版.docx'
    result = parse_file(test_path)
    print(f"书名: {result['title']}")
    print(f"作者: {result['author']}")
    print(f"国籍: {result['author_nationality']}")
    print(f"版本: {result['version']}")
    print(f"章数: {len(result['chapters'])}")
    print(f"节数: {len(result['sections'])}")
    print()
    for sec in result['sections'][:3]:
        print(f"第{sec['section_number']}节: {sec['title']} ({len(sec['content'])}字)")
        print(f"  内容预览: {sec['content'][:80]}...")
        print()
