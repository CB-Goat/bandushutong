"""
调试 Word 文档批注位置
用法：python3 debug_word_comments.py <docx文件路径>
"""
import sys
from docx import Document
from lxml import etree

def debug_comments(file_path):
    doc = Document(file_path)
    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    nsmap = {'w': W_NS}
    
    print("=" * 60)
    print("调试 Word 文档批注位置")
    print("=" * 60)
    
    # 打印所有段落（带索引和样式）
    print("\n[段落列表]:")
    for i, para in enumerate(doc.paragraphs[:15]):
        style = para.style.name if para.style else 'Normal'
        text = para.text[:50] if para.text else '(空)'
        print(f"  [{i}] [{style}] {text}...")
    
    # 提取批注范围
    print("\n[批注范围]:")
    comment_ranges = {}
    for para_idx, para in enumerate(doc.paragraphs):
        char_offset = 0
        for child in para._element:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            
            if tag == 'commentRangeStart':
                cid = child.get(f'{{{W_NS}}}id')
                if cid:
                    comment_ranges[cid] = {
                        'start_para': para_idx,
                        'start_char': char_offset
                    }
                    print(f"  段落{para_idx} char_offset={char_offset}: commentRangeStart id={cid}")
            
            elif tag == 'commentRangeEnd':
                cid = child.get(f'{{{W_NS}}}id')
                if cid and cid in comment_ranges:
                    comment_ranges[cid]['end_para'] = para_idx
                    comment_ranges[cid]['end_char'] = char_offset
                    print(f"  段落{para_idx} char_offset={char_offset}: commentRangeEnd id={cid}")
            
            elif tag == 'r':
                for t in child.findall('.//w:t', nsmap):
                    if t.text:
                        char_offset += len(t.text)
    
    # 提取批注文本
    print("\n[批注内容]:")
    comments = {}
    for rel in doc.part.rels.values():
        reltype = str(rel.reltype)
        if reltype.endswith('/comments') and 'commentsExtended' not in reltype:
            root = etree.fromstring(rel.target_part.blob)
            for comment in root.findall('.//w:comment', nsmap):
                cid = comment.get(f'{{{W_NS}}}id')
                texts = []
                for p in comment.findall('.//w:p', nsmap):
                    for t in p.findall('.//w:t', nsmap):
                        if t.text:
                            texts.append(t.text)
                text = ''.join(texts).strip()
                comments[cid] = text
                print(f"  ID={cid}: {text[:50]}...")
    
    # 显示批注引用的原文
    print("\n[批注引用原文]:")
    for cid, rng in comment_ranges.items():
        para_idx = rng['start_para']
        start_char = rng['start_char']
        end_char = rng['end_char']
        para_text = doc.paragraphs[para_idx].text
        
        print(f"\n  批注 ID={cid}:")
        print(f"    段落={para_idx}, 字符范围=[{start_char}:{end_char}]")
        print(f"    段落全文: \"{para_text[:80]}...\"")
        print(f"    引用原文: \"{para_text[start_char:end_char]}\"")
        if cid in comments:
            print(f"    批注内容: \"{comments[cid][:50]}...\"")

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("用法: python3 debug_word_comments.py <docx文件路径>")
        sys.exit(1)
    debug_comments(sys.argv[1])
